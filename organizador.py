import os
import shutil
import re
import json
import warnings
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm
from transformers import pipeline
from huggingface_hub import login
import torch
from PyPDF2 import PdfReader
from ebooklib import epub
import docx
from difflib import get_close_matches
import unicodedata

# Constantes
login(token="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")

CARPETA_ENTRADA = 'Libros'
CARPETA_SALIDA = 'Libros_Organizados'
LOG_FILE = 'errores_procesamiento.json'
MAX_WORKERS = os.cpu_count()
MAX_PAGES = 10
MAX_PARAGRAPHS_PER_PAGE = 30
MAX_EPUB_ITEMS = 10
MAX_CHARACTERS = 5000  # Limit the text length to prevent tokenizer overflow

# Configuración del dispositivo
device = "cuda" if torch.cuda.is_available() else "cpu"

torch.backends.cudnn.benchmark = True
if device == "cuda":
    torch.cuda.set_per_process_memory_fraction(0.9)
else:
    torch.set_num_threads(os.cpu_count())

qa_pipeline = pipeline(
    "question-answering",
    model="mrm8488/bert-base-spanish-wwm-cased-finetuned-spa-squad2-es",
    tokenizer="mrm8488/bert-base-spanish-wwm-cased-finetuned-spa-squad2-es",
    device=0 if device == "cuda" else -1
)

ner_pipeline = pipeline(
    "ner",
    model="mrm8488/bert-spanish-cased-finetuned-ner",
    tokenizer="mrm8488/bert-spanish-cased-finetuned-ner",
    aggregation_strategy="simple",
    device=0 if device == "cuda" else -1
)

# Obtener la longitud máxima permitida por los modelos
MAX_LENGTH_QA = qa_pipeline.tokenizer.model_max_length
MAX_LENGTH_NER = ner_pipeline.tokenizer.model_max_length

# Usar la longitud máxima mínima entre los dos modelos
MAX_LENGTH = min(MAX_LENGTH_QA, MAX_LENGTH_NER)

known_authors = set()

def clean_text(text):
    # Remove invalid characters and normalize whitespace
    text = text.encode('utf-8', 'ignore').decode('utf-8', 'ignore')
    text = unicodedata.normalize('NFKD', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def normalize_author_name(name):
    name = name.lower()
    name = ''.join(
        c for c in unicodedata.normalize('NFD', name)
        if unicodedata.category(c) != 'Mn'
    )
    name = re.sub(r'[^a-z\s]', '', name)
    name = ' '.join(name.split())
    return name

def get_best_matching_author(name):
    matches = get_close_matches(name, known_authors, cutoff=0.8)
    if matches:
        return matches[0]
    else:
        known_authors.add(name)
        return name

def chunk_text(text, tokenizer, max_length):
    tokens = tokenizer.encode(text, add_special_tokens=False)
    num_tokens = len(tokens)
    for i in range(0, num_tokens, max_length):
        token_chunk = tokens[i:i+max_length]
        decoded_chunk = tokenizer.decode(token_chunk, skip_special_tokens=True)
        yield decoded_chunk

def extract_author_name(text):
    patterns = [
        r'(?i)Autor[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+(?:\s[A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)+)',
        r'(?i)Escrito por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+(?:\s[A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)+)',
        r'(?i)Por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+(?:\s[A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)+)',
        r'(?i)Autor\([^\)]+\):\s*([^\n]+)',
        r'(?i)Autor[^\n]+:\s*([^\n]+)',
        r'(?i)Escrito por[^\n]+:\s*([^\n]+)',
        r'(?i)Por[^\n]+:\s*([^\n]+)',
        r'(?i)Autor[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)',
        r'(?i)Escrito por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)',
        r'(?i)Por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)',
        r'(?i)Por\([^\)]+\):\s*([^\n]+)',
        r'(?i)Escrito por\([^\)]+\):\s*([^\n]+)',
        r'(?i)Traducido por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+(?:\s[A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)+)',
        r'(?i)Traducido por[:\s]+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ]+)',
        r'(?i)Traducido por\([^\)]+\):\s*([^\n]+)',
        r'(?i)Traducido por[^\n]+:\s*([^\n]+)'
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()

    # Use NER to find person entities
    for chunk in chunk_text(text, ner_pipeline.tokenizer, MAX_LENGTH_NER):
        try:
            entities = ner_pipeline(chunk)
            for entity in entities:
                if entity['entity_group'] == 'PER':
                    return entity['word']
        except Exception:
            continue
    return None

def process_pdf(ruta_archivo):
    try:
        lector = PdfReader(ruta_archivo)
        num_paginas = min(MAX_PAGES, len(lector.pages))
        texto = ''
        for num_pagina in range(num_paginas):
            pagina = lector.pages[num_pagina]
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto += texto_pagina + '\n'
        return clean_text(texto)
    except Exception as e:
        return f"Error processing PDF {ruta_archivo}: {e}"

def process_epub(ruta_archivo):
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            libro = epub.read_epub(ruta_archivo)
        texto = ''
        conteo = 0
        for item in libro.get_items():
            if item.get_type() == epub.EpubHtml:
                contenido = item.get_content().decode('utf-8', errors='ignore')
                contenido = re.sub(r'<[^>]+>', '', contenido)
                texto += contenido + '\n'
                conteo += 1
                if conteo >= MAX_EPUB_ITEMS:
                    break
        return clean_text(texto)
    except Exception as e:
        return f"Error processing EPUB {ruta_archivo}: {e}"

def process_docx(ruta_archivo):
    try:
        documento = docx.Document(ruta_archivo)
        texto = ''
        num_parrafos = min(MAX_PAGES * MAX_PARAGRAPHS_PER_PAGE, len(documento.paragraphs))
        for i in range(num_parrafos):
            texto += documento.paragraphs[i].text + '\n'
        return clean_text(texto)
    except Exception as e:
        return f"Error processing DOCX {ruta_archivo}: {e}"

def process_file(args):
    ruta_archivo, ext = args
    if ext == '.pdf':
        result = process_pdf(ruta_archivo)
    elif ext == '.epub':
        result = process_epub(ruta_archivo)
    elif ext == '.docx':
        result = process_docx(ruta_archivo)
    else:
        result = None
    return ruta_archivo, result

def main():
    archivos_error = []
    archivos_no_soportados = []

    if not os.path.exists(CARPETA_SALIDA):
        os.makedirs(CARPETA_SALIDA)

    archivos_para_procesar = []
    for root, _, files in os.walk(CARPETA_ENTRADA):
        for nombre_archivo in files:
            ruta_archivo = os.path.join(root, nombre_archivo)
            if os.path.isfile(ruta_archivo):
                ext = os.path.splitext(ruta_archivo)[1].lower()
                if ext in ['.pdf', '.epub', '.docx']:
                    archivos_para_procesar.append((ruta_archivo, ext))
                else:
                    archivos_no_soportados.append(ruta_archivo)

    textos_para_procesar = []
    rutas_archivos = []
    with ProcessPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(process_file, args): args[0] for args in archivos_para_procesar}
        for future in tqdm(as_completed(futures), total=len(futures), desc="Extrayendo textos de archivos", unit="archivo"):
            ruta_archivo = futures[future]
            try:
                ruta_archivo, result = future.result()
                if result is None or result.startswith("Error processing"):
                    archivos_error.append((ruta_archivo, result))
                else:
                    textos_para_procesar.append(result)
                    rutas_archivos.append(ruta_archivo)
            except Exception as e:
                archivos_error.append((ruta_archivo, str(e)))

    authors = []
    for idx, text in enumerate(tqdm(textos_para_procesar, desc="Extrayendo autores", unit="archivo")):
        author = extract_author_name(text)
        if not author:
            # Limit the context to the first MAX_CHARACTERS characters
            context = text[:MAX_CHARACTERS]
            try:
                # Ensure the context is within tokenizer limits
                context_tokens = qa_pipeline.tokenizer.encode(context, add_special_tokens=False)
                if len(context_tokens) > MAX_LENGTH_QA:
                    context_tokens = context_tokens[:MAX_LENGTH_QA]
                context = qa_pipeline.tokenizer.decode(context_tokens, skip_special_tokens=True)
                qa_input = {'context': context, 'question': '¿Quién es el autor del libro?'}
                output = qa_pipeline(**qa_input)
                author = output.get('answer', None)
            except Exception as e:
                author = None
                archivos_error.append((rutas_archivos[idx], f"Error processing QA: {e}"))
        authors.append(author)

    for idx, author in enumerate(tqdm(authors, desc="Organizando archivos por autor", unit="archivo")):
        ruta_archivo = rutas_archivos[idx]
        nombre_archivo = os.path.basename(ruta_archivo)
        try:
            if not author or author.strip() == '' or author.lower() == 'no answer':
                nombre_autor = 'Autor Desconocido'
            else:
                nombre_autor = normalize_author_name(author)
                nombre_autor = get_best_matching_author(nombre_autor)
                nombre_autor = re.sub(r'[<>:"/\\|?*]', '', nombre_autor)

            carpeta_autor = os.path.join(CARPETA_SALIDA, nombre_autor)

            if not os.path.exists(carpeta_autor):
                os.makedirs(carpeta_autor)

            shutil.copy2(ruta_archivo, os.path.join(carpeta_autor, nombre_archivo))
        except Exception as e:
            archivos_error.append((ruta_archivo, str(e)))

    log_data = {
        "archivos_error": archivos_error,
        "archivos_no_soportados": archivos_no_soportados
    }
    with open(LOG_FILE, 'w', encoding='utf-8') as log_file:
        json.dump(log_data, log_file, indent=4, ensure_ascii=False)

    print("Terminé")

if __name__ == '__main__':
    main()

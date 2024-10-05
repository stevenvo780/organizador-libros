import os
import io
import re
import warnings
from contextlib import redirect_stderr

import fitz
from PyPDF2 import PdfReader
from pdfminer.high_level import extract_text as pdfminer_extract_text
import torch
import easyocr
from PIL import Image
import docx
from ebooklib import epub
import pypandoc
from utils import clean_text, log_error
from file_types import FORMATOS_ARCHIVOS

MAX_PAGES = 10
MAX_PARAGRAPHS_PER_PAGE = 30
MAX_EPUB_ITEMS = 10
MAX_CHARACTERS = 5000
reader = easyocr.Reader(['en'], gpu=torch.cuda.is_available())

def fragment_text(text, max_characters=MAX_CHARACTERS):
    if len(text) <= max_characters:
        return [text]
    fragmentos = [
        text[:max_characters],
        text[len(text) // 2 : len(text) // 2 + max_characters],
        text[-max_characters:],
    ]
    return fragmentos

def extract_metadata_default(ruta_archivo):
    filename = os.path.basename(ruta_archivo)
    return {'author': '', 'title': '', 'filename': filename}

def extract_metadata_pdf(documento, ruta_archivo):
    metadata = documento.metadata
    autor = metadata.get("author", "")
    titulo = metadata.get("title", "")
    filename = os.path.basename(ruta_archivo)
    return {'author': autor, 'title': titulo, 'filename': filename}

def extract_metadata_epub(libro, ruta_archivo):
    titulo = ''
    autor = ''
    titulo_meta = libro.get_metadata('DC', 'title')
    if titulo_meta:
        titulo = titulo_meta[0][0]
    autor_meta = libro.get_metadata('DC', 'creator')
    if autor_meta:
        autor = autor_meta[0][0]
    filename = os.path.basename(ruta_archivo)
    return {'author': autor, 'title': titulo, 'filename': filename}

def extract_metadata_docx(documento, ruta_archivo):
    core_properties = documento.core_properties
    autor = core_properties.author or ''
    titulo = core_properties.title or ''
    filename = os.path.basename(ruta_archivo)
    return {'author': autor, 'title': titulo, 'filename': filename}

def extract_images_from_pdf(documento):
    texto_extraido = ""
    for num_pagina in range(min(MAX_PAGES, len(documento))):
        pagina = documento.load_page(num_pagina)
        imagenes = pagina.get_images(full=True)
        for img in imagenes:
            xref = img[0]
            base_imagen = documento.extract_image(xref)
            imagen_bytes = base_imagen["image"]
            imagen = Image.open(io.BytesIO(imagen_bytes))
            ocr_resultado = reader.readtext(imagen)
            for resultado in ocr_resultado:
                texto_extraido += resultado[1] + "\n"
    return texto_extraido

def process_pdf(ruta_archivo):
    try:
        with io.StringIO() as buf, redirect_stderr(buf):
            texto = pdfminer_extract_text(ruta_archivo, maxpages=MAX_PAGES)
            if texto.strip():
                return clean_text(texto), extract_metadata_default(ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing PDF with pdfminer.six: {e}")

    try:
        documento = fitz.open(ruta_archivo)
        texto = ""
        metadata = extract_metadata_pdf(documento, ruta_archivo)
        for num_pagina in range(min(MAX_PAGES, len(documento))):
            pagina = documento.load_page(num_pagina)
            texto += pagina.get_text() + "\n"
        texto_imagenes = extract_images_from_pdf(documento)
        texto += texto_imagenes
        if texto.strip():
            return clean_text(texto), metadata
    except Exception as e:
        log_error(ruta_archivo, f"Error processing PDF with PyMuPDF: {e}")

    try:
        lector = PdfReader(ruta_archivo)
        num_paginas = min(MAX_PAGES, len(lector.pages))
        texto = ''
        for num_pagina in range(num_paginas):
            pagina = lector.pages[num_pagina]
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto += texto_pagina + '\n'
        if texto.strip():
            return clean_text(texto), extract_metadata_default(ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing PDF with PyPDF2: {e}")

    log_error(ruta_archivo, "No se pudo extraer el texto del PDF.")
    return None, None

def process_epub(ruta_archivo):
    try:
        with io.StringIO() as buf, redirect_stderr(buf):
            libro = epub.read_epub(ruta_archivo)
            texto = ''
            conteo = 0
            for item in libro.get_items():
                if item.get_type() == epub.EpubHtml:
                    contenido = item.get_content().decode('utf-8', errors='ignore')
                    contenido = re.sub(r'<[^>]+>', '', contenido)
                    texto += contenido + '\n'
                    conteo += 1
                    if conteo >= MAX_EPUB_ITEMS:
                        break
        return clean_text(texto), extract_metadata_epub(libro, ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing EPUB: {e}")
        return None, None

def process_docx(ruta_archivo):
    try:
        documento = docx.Document(ruta_archivo)
        texto = ''
        num_parrafos = min(MAX_PAGES * MAX_PARAGRAPHS_PER_PAGE, len(documento.paragraphs))
        for i in range(num_parrafos):
            texto += documento.paragraphs[i].text + '\n'
        return clean_text(texto), extract_metadata_docx(documento, ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing DOCX: {e}")
        return None, None

def process_doc(ruta_archivo):
    try:
        texto = pypandoc.convert_file(ruta_archivo, 'plain', format='doc')
        return clean_text(texto), extract_metadata_default(ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing DOC: {e}")
        return None, None

def process_rtf(ruta_archivo):
    try:
        texto = pypandoc.convert_file(ruta_archivo, 'plain', format='rtf')
        return clean_text(texto), extract_metadata_default(ruta_archivo)
    except Exception as e:
        log_error(ruta_archivo, f"Error processing RTF: {e}")
        return None, None

def process_file(ruta_archivo, ext):
    if ext in FORMATOS_ARCHIVOS['pdf']:
        return process_pdf(ruta_archivo)
    elif ext in FORMATOS_ARCHIVOS['epub']:
        return process_epub(ruta_archivo)
    elif ext in FORMATOS_ARCHIVOS['docx']:
        return process_docx(ruta_archivo)
    elif ext in FORMATOS_ARCHIVOS['doc']:
        return process_doc(ruta_archivo)
    elif ext in FORMATOS_ARCHIVOS['rtf']:
        return process_rtf(ruta_archivo)
    else:
        log_error(ruta_archivo, f"Unsupported file type: {ext}")
        return None, None
